import re
from docx import Document
from converters.utils import (
    setup_database, parse_parts, parse_footnote_markers, 
    get_part, get_page_reset_ids, clean_content_text, clean_toc_title,
    wrap_direction, parse_skip_ids
)

def process_docx(docx_path, db_path, book_id="00000", page_marker=None, 
                 footnote_markers_str=None, parts_str=None, 
                 collapse_newlines=False, clean_toc=False, fix_cap=False,
                 detect_dir=False, skip_ids_str=None, max_len_skip=0,
                 progress_callback=None, footnote_separator=None):
    
    footnote_markers = parse_footnote_markers(footnote_markers_str)
    part_boundaries = parse_parts(parts_str)
    page_reset_ids  = get_page_reset_ids(part_boundaries)
    skip_ids_set    = parse_skip_ids(skip_ids_str)

    doc = Document(docx_path)
    conn, cur, table_b, table_t = setup_database(db_path, book_id)

    global_id = 1
    source_page_counter = 1
    current_page_number = 1
    pending_headings = []
    current_page_text = []

    table_b_inserts = []
    table_t_inserts = []

    def flush_page():
        nonlocal global_id, source_page_counter, current_page_number, current_page_text, table_b_inserts, table_t_inserts
        raw_text = "\n".join(current_page_text).strip()
        text = clean_content_text(raw_text, collapse_newlines=collapse_newlines, footnote_separator=footnote_separator)
        text = wrap_direction(text, enable=detect_dir)
        
        if not text:
            return
        
        # SKIP LOGIC
        should_skip = False
        if source_page_counter in skip_ids_set:
            should_skip = True
        if max_len_skip > 0 and len(text) > max_len_skip:
            should_skip = True
            
        if should_skip:
            source_page_counter += 1
            current_page_text = []
            # We DON'T clear pending_headings here, 
            # so they will point to the next non-skipped page.
            return

        if global_id in page_reset_ids:
            current_page_number = 1
            
        part = get_part(global_id, part_boundaries)
        
        table_b_inserts.append((text, part, global_id, current_page_number))
        
        for h_text, h_lvl in pending_headings:
            cleaned_h = clean_toc_title(h_text, remove_numbers=clean_toc, fix_capitalization=fix_cap)
            if cleaned_h:
                table_t_inserts.append((cleaned_h, h_lvl, 0, global_id))
        pending_headings.clear()
        
        global_id += 1
        source_page_counter += 1
        current_page_number += 1
        current_page_text = []

    total_paras = len(doc.paragraphs)

    # ⚡ Bolt: Throttling Streamlit UI updates to max ~100 times per document
    # Updating UI thousands of times blocks the main thread and drastically slows down conversion.
    update_step = max(1, total_paras // 100)

    # ⚡ Bolt: Cache styles to avoid slow O(N) lookup from para.style.name in python-docx
    style_map = {style.style_id: style.name for style in doc.styles}

    # ⚡ Bolt: Pre-compile heading regex to avoid per-paragraph compilation overhead
    heading_pattern = re.compile(r'Heading (\d+)')

    for i, para in enumerate(doc.paragraphs):
        if progress_callback and (i % update_step == 0 or i == total_paras - 1):
            progress_callback(i / total_paras, f"Memproses paragraf {i+1}/{total_paras}")
            
        # Fast dictionary lookup using internal style ID instead of para.style.name (~20x faster)
        style_id = para._p.pPr.pStyle.val if para._p.pPr is not None and para._p.pPr.pStyle is not None else None
        style_name = style_map.get(style_id, "Normal")
        text = para.text.strip()
        
        if not text and not heading_pattern.match(style_name):
            continue

        heading_match = heading_pattern.match(style_name)
        if heading_match:
            if current_page_text:
                flush_page()
            
            if text:
                lvl = int(heading_match.group(1))
                pending_headings.append((text, lvl))
        
        if text:
            if page_marker and page_marker in para.text:
                parts = para.text.split(page_marker)
                for i, p in enumerate(parts):
                    if p.strip():
                        current_page_text.append(p.strip())
                    if i < len(parts) - 1:
                        flush_page()
            else:
                current_page_text.append(text)

    flush_page()

    if table_b_inserts:
        cur.executemany(
            f"INSERT INTO {table_b} (nass, part, id, page) VALUES (?, ?, ?, ?)",
            table_b_inserts
        )
    if table_t_inserts:
        cur.executemany(
            f"INSERT INTO {table_t} (tit, lvl, sub, id) VALUES (?, ?, ?, ?)",
            table_t_inserts
        )

    conn.commit()
    conn.close()
    return True
